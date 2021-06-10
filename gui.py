# to get started read this https://realpython.com/python-gui-tkinter/
# watch this  https://www.youtube.com/watch?v=YXPyB4XeYLA
# done, you know tkinter now!
from tkinter import *
from tkinter import filedialog
from tkinter import ttk
from ttkthemes import ThemedTk
from PIL import ImageTk, Image
from docx import *
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from main import prepare_img_keras, predict_class
from data_base import save_in_db, query_by_name, c
# from encryption import encrypt, decrypt
import  re

class App:
    def __init__(self):
        print('App created..')

    def gmtry(self, app_width, app_height):
        self.app_width = app_width
        self.app_height = app_height
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.winfo_screenwidth()/2 - app_width/2)
        positionDown = int(self.winfo_screenheight()/2 - app_height/2)
        # Positions the window in the center of the page.
        self.geometry(f"{app_width}x{app_height}+{positionRight}+{positionDown}")

    def btn(self, frame, text, command ):
        return ttk.Button(frame, text = text,width=17, command=command)

    def img():
        pass


class Main_window(ThemedTk, App):
    def __init__(self):
        super().__init__()
        global m_background
        m_background = ImageTk.PhotoImage(Image.open('./assets/bk_main.png'))
        Label(self, image =m_background).pack(expand=True)
        self.body = ttk.Frame(self)
        self.body.place(relx=0.5,rely=0.5,anchor=CENTER) # center all app content to the center of the body
        self.resizable(width = True, height = True) # Allow Window to be resizable
        self.title('Doctorizer')
        self.image_path = None
        self.result_frame = None
        self.result_label = None
        self.result = ''
        self.iconbitmap('./assets/logo.ico')
        # self.config(bg="white")
        #left and right frames
        self.l_upp_frame = ttk.Frame(self.body)
        self.l_upp_frame.place(x=0,y=0)

        self.left_frame = ttk.Frame(self.body)
        self.left_frame.grid(row=1, column=0, padx=50, pady=10)

        self.right_frame = ttk.LabelFrame(self.body, text='X-ray image', width=505,height=510)
        self.right_frame.grid(row=1, column=1, padx=50,pady=10)
        # global b_background
        # b_background = ImageTk.PhotoImage(Image.open('./assets/body_bk.png'))
        # Label(self.right_frame, image =b_background).place(x=0,y=0)
        # main screen widgets
        # Label(self, text ="X-ray image analyzer", font=10).place(relx=0.5,rely=0.04, anchor=CENTER)
        global db_image
        db_image = Image.open('./assets/db.png')
        db_image = db_image.resize((25, 25))
        db_image = ImageTk.PhotoImage(db_image)
        Button(self.l_upp_frame, image = db_image,highlightthickness = 0, bd = 0, command=show_database_w).pack(padx=5,pady=5)

        self.btn(self.left_frame, text ="upload image", command=self.upload_image).grid(row=1, column=0,padx=50, pady=10)
        self.btn(self.left_frame, text ="clear workspace",command = lambda: self.clear_workspace(panel)).grid(row=2, column=0, pady=10)
        self.btn(self.left_frame, text ="classify", command=self.classify).grid(row=3, column=0, pady=10)
        self.btn(self.left_frame, text ="automated report", command=show_patient_dw).grid(row=4, column=0, pady=10)

    #functions
    def upload_image(self):
        # Select the Imagename from a folder
        try:
            self.image_path = filedialog.askopenfilename(title ='"pen')
        except:
            print('Failed To Open the File')
        if self.image_path:
            global img
            # opens the image
            img = Image.open(self.image_path)
            # resize the image and rootly a high-quality down sampling filter
            img = img.resize((500, 500), Image.ANTIALIAS)
            # PhotoImage class is used to add image to widgets, icons etc
            img = ImageTk.PhotoImage(img)
            # create a label
            global panel
            panel = ttk.Label(self.right_frame, image = img)
            # set the image as img
            panel.grid(row=0, column=0)

    def clear_workspace(self, space):
        print('clear_workspace')
        if space:
            space.grid_forget()
        if self.result_frame:
            print('cleared result')
            self.result_label.grid_forget()
            self.result_frame.grid_forget()

    def classify(self):
        print('classify')
        self.result = ''
        if self.image_path:
            prepared_img = prepare_img_keras(self.image_path)
            self.result = predict_class(prepared_img)
            self.show_classification_result()

    def show_classification_result(self):
        self.result_frame = ttk.LabelFrame(self.body, text='Classification result', width=505,height=100)
        self.result_frame.grid(row=2, column=1, padx=50,pady=10)
        self.result_label = ttk.Label(self.result_frame, text=self.result)
        self.result_label.grid(padx=50,pady=5)

class Window(Toplevel, App):
    def __init__(self, parent):
        super().__init__(parent)
        global wn_background
        wn_background = ImageTk.PhotoImage(Image.open('./assets/bk.png'))
        Label(self, image = wn_background).pack(expand=True)
        self.body = ttk.Frame(self)
        # center all app content to the center of the body
        self.body.place(relx=0.5,rely=0.5,anchor=CENTER)
        # Allow Window to be resizable
        self.resizable(width=False, height=False)
        self.title('Doctorizer')
        self.iconbitmap('./assets/logo.ico')
        # self.config(bg="#2d2d2d")


class Patient_data_window(Window):
    def __init__(self, parent):
        super().__init__(parent)
        #window settings
        self.title('patient data')

        self.normal = """The chest X-ray image is a normal chest X-ray showing no obvious abnormality.
The patient must continue to protect himself and others by adhering to preventive measures, as announced by the World Health Organization."""

        self.covid = """The chest X-ray image shows abnormalities like consolidation, ground-glass opacity, nodular shadowing, and increase in lung density, the most likely cause for this is pneumonia caused by covid 19. A follow-up chest X-ray should be performed in 4-6 weeks following the commencement of treatment to ensure the resolution of pneumonia.
The patient must adhere to the health protocol followed by the World Health Organization and adhere to the home quarantine for a period of 14 days.
The patient needs regular pulse oximetry to monitor O2 saturation if O2 saturation< than 90%, the patient needs for O2 supplement.
The patient may take certain drugs as needed for fever, discomfort, and increase immunity. These include Levofloxacin Tablet 500mg *1 for 7 days, Azithromycin 250mg 2*1 for 6 days, Paracetamol 500 mg*3, and Vitamins complex."""

        self.pneumonia_bacteria = """This chest X-ray shows an abnormality, the most likely cause for this is bacteria pneumonia. The patient should be assessed using the CURB 65 clinical scoring system and treated with appropriate antibiotics. A follow-up chest X-ray should be performed in 4-6 weeks following the commencement of treatment to ensure the resolution of pneumonia.
The patient may take these as needed for fever and discomfort. These include drugs such as aspirin, ibuprofen (Advil, Motrin IB, others) and acetaminophen (Tylenol, others)."""

        self.pneumonia_viral = """The chest X-ray image shows abnormalities like bilateral perihilar peribranchial thickening, interstitial infiltrates and small-caliber airways have areas of atelectasis or air trapping,  the most likely cause for this is viral pneumonia. A follow-up chest X-ray should be performed in 4-6 weeks following the commencement of treatment to ensure the resolution of pneumonia.
The patient should get sufficient rest and stay hydrated by drinking plenty of fluids.
The patient needs regular pulse oximetry to monitor O2 saturation if O2 saturation< than 90%, the patient needs for O2 supplement.
The patient may take certain drugs as needed for fever, discomfort, and increase immunity. These include Levofloxacin Tablet 500mg *1 for 7 days, Azithromycin 250mg 2*1 for 6 days, Paracetamol 500 mg*3, and Vitamins complex."""

        ttk.Label(self.body, text='Patient Name').grid(row=0, column=0)
        entry0 = ttk.Entry(self.body)
        entry0.grid(row=0, column=1,padx=5, pady=10)


        ttk.Label(self.body, text='Patient Age').grid(row=1, column=0)
        entry1 = ttk.Entry(self.body)
        entry1.grid(row=1, column=1, padx=5,pady=10)


        ttk.Label(self.body, text='Chronic diseases').grid(row=2, column=0)
        entry2 = ttk.Entry(self.body)
        entry2.grid(row=2, column=1, padx=5,pady=10)

        ttk.Label(self.body, text='gender').grid(row=3, column=0)
        entry3 = ttk.Entry(self.body)
        entry3.grid(row=3, column=1, padx=5,pady=10)


        ttk.Label(self.body, text='blood type').grid(row=4, column=0)
        entry4 = ttk.Entry(self.body)
        entry4.grid(row=4, column=1, padx=5,pady=10)

        ttk.Label(self.body, text='date').grid(row=5, column=0)
        entry5 = ttk.Entry(self.body)
        entry5.grid(row=5, column=1, padx=5,pady=10)

        self.btn(self.body, text='generate report', command = lambda: [self.generate_report(
        name = entry0.get(),
        age = entry1.get(),
        ill = entry2.get(),
        gender = entry3.get(),
        blood_type = entry4.get(),
        date = entry5.get(),
        result = parent.result,
        ),self.destroy()]
        ).grid(row=7,column=2, padx=10,pady=15)


        self.btn(self.body, text='Save to database', command = lambda: save_in_db(
        full_name = entry0.get(),
        age = entry1.get(),
        cronic_diseas = entry2.get(),
        gender = entry3.get(),
        blood_type = entry4.get(),
        image_date = entry5.get(),
        result = parent.result,
        img_data = str(ImageTk.PhotoImage(Image.open(root.image_path)))
        )).grid(row=7,column=1, padx=10,pady=15)

    def generate_report(self,name,age,ill, result,gender,blood_type,date): # report: this function ******************************
        document = Document()
        # from here
        my_image = document.add_picture('assets/2632 [Converted]+1.png', width=Inches(1.25))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_heading('Doctorizer automated report', 0)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_heading('Patient info:', 1)

        info = document.add_paragraph().add_run(f'Patient Name: {name} \nPatient Age: {age} \nChronic diseases: {ill} \ngender: {gender} \nBlood: {blood_type} \ndate: {date}')
        info.font.size = Pt(13)
        info.font.name = 'Roboto'
        info.font.color.rgb = RGBColor(55, 55, 95)
        document.add_heading('Case:',1)
        document.add_paragraph().add_run(f'{result}').bold = True

        # to here
        print(result)
        if result == 'No illness detected':
            case = self.normal
        elif result == 'Pneumonia | Bacteria':
            case = self.pneumonia_bacteria
        elif result == 'Pneumonia | COVID-19 Positive':
            case = self.covid
        else:
            case = f'mmm just a {result}'

        document.add_heading('State And Recomendation', 2)
        para = document.add_paragraph().add_run(case)
        para.font.size = Pt(11)
        para.font.name = 'Roboto'
        document.add_picture('assets/Untitled-5.png', width=Inches(0.85))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print(result)
        self.save_report(document)

    def save_report(self, document):
        filepath = filedialog.asksaveasfilename(initialfile = 'report.docx',
        defaultextension=".docx",
        filetypes=[("All Files","*.*"),
        ("Word Documents","*.docx"),
        ])
        try:
            document.save(filepath)
        except:
            raise Exception('Error: File not saved')

        self.open_report(filepath)

    def open_report(self,filepath):
        import subprocess, os, platform
        if platform.system() == 'Darwin':       # macOS
            subprocess.call(('open', filepath))
        elif platform.system() == 'Windows':    # Windows
            os.startfile(filepath)
        else:                                   # linux variants
            subprocess.call(('xdg-open', filepath))


class database_window(Window):
    def __init__(self, parent):
        super().__init__(parent)
        #window settings
        global wn_background
        wn_background = ImageTk.PhotoImage(Image.open('./assets/bk_db.png'))
        Label(self, image = wn_background).place(x=0,y=0)
        self.title('database')

        t = ttk.Treeview(self)
        t['columns'] = ('name','age','chronic diseases' , 'gender','blood_type','date','case')
        t['show']='headings'
        t.column('name',width=100,minwidth=100, anchor=CENTER)
        t.column('age',width=50,minwidth=50, anchor=CENTER )
        t.column('chronic diseases',width=100,minwidth=100, anchor=CENTER)
        t.column('gender',width=70,minwidth=70, anchor=CENTER)
        t.column('blood_type',width=70,minwidth=70, anchor=CENTER)
        t.column('date',width=70,minwidth=70, anchor=CENTER)
        t.column('case',width=200,minwidth=200, anchor=CENTER)


        t.heading('name', anchor=CENTER, text='name')
        t.heading('age', anchor=CENTER, text='age')
        t.heading('chronic diseases', anchor=CENTER, text='chronic diseases')
        t.heading('gender', anchor=CENTER, text='gender')
        t.heading('blood_type', anchor=CENTER, text='blood type')
        t.heading('date', anchor=CENTER, text='date')
        t.heading('case', anchor=CENTER, text='case')

        i=0
        for row in c.execute("SELECT * ,oid FROM information"):
            t.insert('',i,text='',values=(row[0], row[1], row[2], row[3], row[4], row[5], row[6]))
            i += 1
        t.place(x=10,y=20)



if __name__ == '__main__':
    def show_patient_dw():
        patient_dw = Patient_data_window(parent=root)
        patient_dw.gmtry(600,400)

    def show_database_w():
        patient_dw = database_window(parent=root)
        patient_dw.gmtry(800,533)
    #create root instance
    root = Main_window()
    print(root.get_themes())
    root.set_theme('xpnative') #xpnative itft1
    root.gmtry(1100,600)
    root.mainloop()
